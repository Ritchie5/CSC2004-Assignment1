from PIL import Image
from docx import Document
import numpy as np
import base64
import os

class LSB:
    def __init__(self, secret_fileformat):
        self.secret_fileformat = secret_fileformat

    def set_bits(self, bits):
        self.bits = int(bits)
        if not 0 <= self.bits <= 5:
            raise ValueError('[!] Number of bits needs to be between 0 - 5.')
    
    def get_object(self, file, itype):
        try:
            extension = os.path.splitext(file)[-1].lower()
            if extension == ".png":
                data = Image.open(file)
            elif extension == ".txt":
                with open(file, 'rb') as data:
                    data = data.read()
            elif extension == ".docx":
                doc = Document(file)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                text ='\n'.join(fullText)
                data = text.encode('utf-8')
        except IOError:
            print('[!] {} file could not be opened.'.format(itype.title()))
        return extension, data
    
    # Image for now!!
    def save_coverimage(self, data, outfile):
        try:
            data.save(outfile)
        except IOError as e:
            raise IOError('[!] {} file could not be written.'.format(outfile) + '\n[!] {}'.format(e))
        except Exception as e:
            raise Exception('[!] Unable to save file.' + '\n[!] {}'.format(e))

    def to_binary(self, data):
        # convert 'data' to binary format as string
        if isinstance(data, str):
            return ''.join([format(ord(i), "08b") for i in data])
        elif isinstance(data, bytes) or isinstance(data,np.ndarray):
            return [format(i, "08b") for i in data]
        elif isinstance(data, int) or isinstance(data,np.uint8):
            return ''.join(data, "08b")
        else:
            raise TypeError("Type not supported.")

    def to_string(self, data):
            return "".join([chr(int(data[i:i+8],2)) for i in range(0,len(data),8)])
    
    def modify_bit(self, selected_bits, pixel, payload_bit):
        for n in range(0, selected_bits+1):
            mask = 1 << n
            result = (pixel & ~mask) | (payload_bit << n)
            n+=1
        return result

    def write_file(self, extension, text):
        if extension == ".txt":
            with open('secret_message' + extension, 'w') as f:
                f.write(text)
        elif extension == ".docx":
            document = Document()
            document.add_paragraph(text)
            document.save('secret_message' + extension)
        
class Encode(LSB):
    def __init__(self, cover, secret, bits, outfile):
        print('[*] LSB Encoding with bits = {}'.format(bits))
        self.set_bits(bits)
        self.outfile = outfile
        self.cover = self.get_object(cover, 'cover')
        self.secret = self.get_object(secret, 'secret')
        self.encode_object()

    # b64 -> binary
    def to_base64(self):
        encoded_string = base64.b64encode(self.secret[1])
        encoded_extension = base64.b64encode(self.secret[0].encode('utf-8'))
        result = self.to_binary(encoded_string.decode('utf-8'))
        result += self.to_binary('#####') # add stopping criteria
        result += self.to_binary(encoded_extension.decode('utf-8'))
        result += self.to_binary('#####') # add stopping criteria
        return result

    def encode_object(self):
        data_index = 0
        binary_secret_data = self.to_base64()
        with self.cover[1] as img:
            width, height = img.size
            for x in range(0, width):
                for y in range(0, height):
                    pixel = list(img.getpixel((x, y)))
                    #RGB
                    for n in range(0, 3):
                        if(data_index < len(binary_secret_data)):
                            #pixel[n] = self.modify_bit(self.bits, pixel[n], int(binary_secret_data[data_index]))
                            pixel[n] = pixel[n] & ~1 | int(binary_secret_data[data_index])
                            data_index += 1
                    img.putpixel((x, y), tuple(pixel))
            
            self.save_coverimage(img, self.outfile)

class Decode(LSB):
    def __init__(self, steg, bits):
        print('[*] LSB Decoding with bits = {}'.format(bits))
        self.set_bits(bits)
        self.steg = self.get_object(steg, 'steg')
        self.decode_object()
    
    # binary -> b64
    def from_base64(self, data):
        result = self.to_string(data)
        result = result.split("#####")
        message = base64.b64decode(result[0])
        extension = base64.b64decode(result[1])
        return message, extension

    def decode_object(self):
        extracted_bin = []
        with self.steg[1] as img:
            width, height = img.size
            for x in range(0, width):
                for y in range(0, height):
                    pixel = list(img.getpixel((x, y)))
                    for n in range(0, 3):
                        extracted_bin.append(pixel[n] & 1)
        
        data = "".join([str(x) for x in extracted_bin])
        message = self.from_base64(data)[0].decode("utf-8")
        extension = self.from_base64(data)[1].decode("utf-8")
        print(message)
        self.write_file(extension, message)

def main():
    
    payload = input("Enter payload file: ")
    image = input("Enter image file: ")
    steg_image = input("Enter steg image name: ")
    number = input("Enter the bit: ")

    Encode(image, payload, number, steg_image)
    Decode(steg_image, number)

if __name__ == '__main__':
    main()